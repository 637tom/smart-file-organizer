import os
from pathlib import Path
from google import genai
from dotenv import load_dotenv
from PIL import Image
from pypdf import PdfReader
from docx import Document

load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
client = genai.Client(api_key=api_key)

def describe_file(file_path : Path):
    suf = file_path.suffix.lower()
    contents = []

    # Optimized promt for gemini 2.0 flash
    instruction = """ROLE: Senior Content Analyst.
TASK: Analyze the file and provide a concise summary of its content.

RULES (STRICT):
1. NO HEADERS: Do NOT use "Summary:", "Description:", or markdown titles (#).
2. NO PREAMBLE: Start directly with the description. Do NOT use phrases like "This file contains..." or "Here is the analysis...".
3. LANGUAGE: Strictly English.
4. LENGTH: Max 60 words.
5. STYLE: Professional, technical, and direct.

SPECIFIC GUIDELINES:
- IMAGES: Focus on key objects, text, and context.
- DOCUMENTS: Identify type, dates, entities, and main subject.
- CODE: Explain main logic and purpose.
- UNKNOWN: If no content provided (only filename), state what the file likely represents based on its name/extension."""
    contents.append(instruction)
    contents.append(f"Filename: {file_path.name}")
    try:
        if suf == ".jpg" or suf == ".png":
            im = Image.open(file_path)
            im.thumbnail((512, 512))
            contents.append(im)
        elif suf == ".pdf":
            reader = PdfReader(file_path)
            text = reader.pages[0].extract_text()[:3000]
            contents.append(text)
        elif suf == ".docx":
            d = Document(file_path)
            text = "\n".join([p.text for p in d.paragraphs])
            text = text[:3000]
            contents.append(text)
        elif suf in ['.txt', '.py', '.md', '.json', '.js', '.html', '.css', '.cpp', '.cs', '.java', '.rb', '.php', '.go', '.rs', '.swift', '.kt', '.kts', '.ts', '.tsx', '.jsx', '.scss', '.sass', '.less', '.styl', '.stylus', '.sass', '.less', '.styl', '.stylus', '.in', '.out']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                contents.append(f.read(3000))
        else:
            contents.append("No content available for this file")
    except Exception as e:
        print(f"ANALYSIS ERROR for {file_path.name}: {e}")
        return "UNCATEGORIZED"
    try:
        response = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=contents
        )
        return response.text.strip().lower()
    except Exception as e:
        print(f"ANALYSIS ERROR for {file_path.name}: {e}")
        return "UNCATEGORIZED"

def classify_files(files : list, existing_categories : list):
    if len(existing_categories) > 0:
        existing_categories = ", ".join(existing_categories)
    else:
        existing_categories = "None you have to create new categories"

    # Optimized promt for gemini 2.0 flash
    instruction = f"""ROLE: Senior Content Analyst & Organizer.
TASK: Analyze exactly {len(files)} files and categorize them.

CONTEXT - EXISTING CATEGORIES: [{existing_categories}]

BEST FIT PROTOCOL (Follow strictly):
1. CONTENT FIRST: Analyze content, not just format.
2. SPECIFIC OVER GENERIC: Avoid 'MISC' or 'SCREENSHOTS'. Use topics like 'MATH', 'FINANCE', 'CODING'.
3. REUSE vs. CREATE: Use existing categories if they fit; create new ones only for distinct new topics.

RULES (MANDATORY):
1. OUTPUT FORMAT: Exactly {len(files)} pairs in format: CATEGORY|DESCRIPTION
2. SEPARATOR: Use '^%^' between pairs. Do NOT put '^%^' at the end.
3. CATEGORY NAMES: Single word, UPPERCASE.
4. DESCRIPTION: Max 10 words summary.
5. NO PREAMBLE: Do NOT write "Raw text:", "Here is the list", "Output:", or anything else. 
6. NO MARKDOWN: Do not use code blocks (```). 
7. START IMMEDIATELY: Your response must start with the first letter of the first category.

EXAMPLE BEHAVIOR:
INPUT: file1.jpg (math), file2.py (code)
OUTPUT: MATH|Linear algebra matrices^%^PYTHON|Sorting algorithm script"""


    contents = [instruction]
    for i, file_path in enumerate(files):
        suf = file_path.suffix.lower()
        contents.append(f"--- FILE {i+1}: {file_path.name} ---")
        try:
            if suf == ".jpg" or suf == ".png":
                im = Image.open(file_path)
                im.thumbnail((512, 512))
                contents.append(im)
            elif suf == ".pdf":
                reader = PdfReader(file_path)
                text = reader.pages[0].extract_text()[:1500]
                contents.append(text)
            elif suf == ".docx":
                d = Document(file_path)
                text = "\n".join([p.text for p in d.paragraphs])
                text = text[:1500]
                contents.append(text)
            elif suf in ['.txt', '.py', '.md', '.json', '.js', '.html', '.css', '.cpp', '.cs', '.java', '.rb', '.php', '.go', '.rs', '.swift', '.kt', '.kts', '.ts', '.tsx', '.jsx', '.scss', '.sass', '.less', '.styl', '.stylus', '.sass', '.less', '.styl', '.stylus', '.in', '.out']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    contents.append(f.read(1500))
            else:
                contents.append("No content available for this file")
        except Exception as e:
            print(f"ANALYSIS ERROR for {file_path.name}: {e}")
            contents.append("UNCATEGORIZED")
    try:
        response = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=contents
        )
        return response.text.strip().lower()

    except Exception as e:
        print(f"ANALYSIS ERROR for {file_path.name}: {e}")
        return "UNCATEGORIZED"

def find_best_file(files : list, query : str):
    if(len(files) == 1):
        return files[0]
    instructions = """ROLE: Senior Search & Content Specialist.
TASK: Compare the provided files and identify the single BEST match for the search query.

STRICT RULES:
1. OUTPUT: Return ONLY the exact file path of the winning file. 
2. NO CHATTER: No headers, no "The best file is:", no markdown, no quotes. Just the raw path.
3. CONTENT-DRIVEN: Use the provided text or image content to determine the match, not just the filename.
4. NO MATCH: If no file matches the query well, return the path of the first file.

The path you return must be identical to one of the paths provided in the "FILE X:" headers."""
    contents = [instructions]
    contents.append(f"Query: {query}")
    for i, file_path in enumerate(files):
        suf = file_path.suffix.lower()
        contents.append(f"--- FILE {i+1}: {file_path} ---")
        try:
            if suf == ".jpg" or suf == ".png":
                im = Image.open(file_path)
                im.thumbnail((512, 512))
                contents.append(im)
            elif suf == ".pdf":
                reader = PdfReader(file_path)
                text = reader.pages[0].extract_text()[:1500]
                contents.append(text)
            elif suf == ".docx":
                d = Document(file_path)
                text = "\n".join([p.text for p in d.paragraphs])
                text = text[:1500]
                contents.append(text)
            elif suf in ['.txt', '.py', '.md', '.json', '.js', '.html', '.css', '.cpp', '.cs', '.java', '.rb', '.php', '.go', '.rs', '.swift', '.kt', '.kts', '.ts', '.tsx', '.jsx', '.scss', '.sass', '.less', '.styl', '.stylus', '.sass', '.less', '.styl', '.stylus', '.in', '.out']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    contents.append(f.read(1500))
            else:
                contents.append("No content available for this file")
        except Exception as e:
            print(f"ANALYSIS ERROR for file {file_path.name}: {e}")
            return None
    try:
        response = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=contents
        )
        return Path(response.text.strip())

    except Exception as e:
        print(f"ANALYSIS ERROR {e}")
        return None
    